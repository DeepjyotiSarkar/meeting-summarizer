"""Generates a polished, shareable meeting-minutes Word document."""
from io import BytesIO
from docx import Document
from docx.shared import Pt

from models import Meeting


def build_minutes_docx(meeting: Meeting) -> BytesIO:
    doc = Document()

    doc.add_heading(meeting.title or "Meeting Minutes", level=0)

    meta = doc.add_paragraph()
    meta.add_run(f"Status: {meeting.status}").italic = True
    if meeting.sentiment:
        meta.add_run(f"   |   Sentiment: {meeting.sentiment}").italic = True
    if meeting.health_score is not None:
        meta.add_run(f"   |   Meeting Health Score: {meeting.health_score}/100").italic = True

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(meeting.summary_text or "No summary generated.")

    doc.add_heading("Key Decisions", level=1)
    for d in (meeting.key_decisions or []):
        doc.add_paragraph(d, style="List Bullet")
    if not meeting.key_decisions:
        doc.add_paragraph("No decisions recorded.")

    doc.add_heading("Action Items", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Task", "Owner", "Deadline", "Priority"
    for item in meeting.action_items:
        row = table.add_row().cells
        row[0].text = item.description
        row[1].text = item.owner or "-"
        row[2].text = item.deadline or "-"
        row[3].text = item.priority

    doc.add_heading("Open Questions / Risks", level=1)
    for r in (meeting.risks_or_open_questions or []):
        doc.add_paragraph(r, style="List Bullet")
    if not meeting.risks_or_open_questions:
        doc.add_paragraph("None noted.")

    if meeting.follow_up_email_draft:
        doc.add_heading("Suggested Follow-up Email", level=1)
        p = doc.add_paragraph(meeting.follow_up_email_draft)
        for run in p.runs:
            run.font.size = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
